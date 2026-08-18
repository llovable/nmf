#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""manuscript_draft_ko.md -> .docx 변환기 (이 원고 전용 경량 변환)."""

import re
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = "/home/dyan/nmf/mochi_code/paper/manuscript_draft_260818.md"
DST = "/home/dyan/nmf/mochi_code/paper/manuscript_draft_260818.docx"

# ---- LaTeX -> 읽기 쉬운 유니코드 텍스트(근사) ----
LATEX_MAP = [
    (r"\\times", "×"), (r"\\odot", "⊙"), (r"\\to", "→"), (r"\\in", "∈"),
    (r"\\sum", "Σ"), (r"\\prod", "Π"), (r"\\approx", "≈"), (r"\\cdot", "·"),
    (r"\\nabla", "∇"), (r"\\alpha", "α"), (r"\\beta", "β"), (r"\\lambda", "λ"),
    (r"\\sigma", "σ"), (r"\\mu", "μ"), (r"\\theta", "θ"), (r"\\bullet", "•"),
    (r"\\top", "ᵀ"), (r"\\infty", "∞"), (r"\\leq", "≤"), (r"\\geq", "≥"),
    (r"\\quad", "  "), (r"\\,", " "), (r"\\;", " "), (r"\\!", ""),
    (r"\\big", ""), (r"\\Big", ""), (r"\\left", ""), (r"\\right", ""),
    (r"\\mathbb\{E\}", "E"), (r"\\mathbb\{R\}", "ℝ"),
    (r"\\hat\s*U", "Û"), (r"\\hat\{U\}", "Û"), (r"\\hat\s*X", "X̂"), (r"\\hat\{X\}", "X̂"),
    (r"\\hat\{x\}", "x̂"), (r"\\tilde\s*X", "X̃"), (r"\\tilde\{X\}", "X̃"),
    (r"\\lVert", "‖"), (r"\\rVert", "‖"), (r"\\lVert", "‖"),
    (r"\\mathcal\{L\}", "L"),
    (r"\\text\{([^}]*)\}", r"\1"),
    (r"\\\\", " | "),
]


def latex_to_text(s: str) -> str:
    for pat, rep in LATEX_MAP:
        s = re.sub(pat, rep, s)
    # 분수/위첨자/아래첨자 정리(근사)
    s = re.sub(r"\^\{([^}]*)\}", r"^(\1)", s)
    s = re.sub(r"_\{([^}]*)\}", r"_(\1)", s)
    s = s.replace("{", "").replace("}", "")
    s = re.sub(r"\s+", " ", s).strip()
    return s


def strip_inline_math(text: str) -> str:
    # \( ... \) 인라인 수식
    def repl(m):
        return latex_to_text(m.group(1))
    text = re.sub(r"\\\((.*?)\\\)", repl, text)
    return text


def add_runs_with_bold(par, text):
    """**bold** 처리하여 run 추가."""
    text = strip_inline_math(text)
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for p in parts:
        if not p:
            continue
        if p.startswith("**") and p.endswith("**"):
            r = par.add_run(p[2:-2])
            r.bold = True
        else:
            par.add_run(p)


def main():
    src = sys.argv[1] if len(sys.argv) > 1 else SRC
    dst = sys.argv[2] if len(sys.argv) > 2 else DST
    with open(src, encoding="utf-8") as f:
        lines = f.read().split("\n")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # 수평선/빈 줄
        if stripped in ("---", "***", "___"):
            i += 1
            continue
        if stripped == "":
            i += 1
            continue

        # 블록 수식 \[ ... \]  (여러 줄 가능)
        if stripped.startswith(r"\["):
            buf = [stripped]
            while r"\]" not in buf[-1] and i + 1 < n:
                i += 1
                buf.append(lines[i].strip())
            raw = " ".join(buf).replace(r"\[", "").replace(r"\]", "")
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(latex_to_text(raw))
            r.italic = True
            i += 1
            continue

        # 표 (마크다운 파이프 테이블)
        if stripped.startswith("|") and "|" in stripped[1:]:
            table_lines = []
            while i < n and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            # 구분선(---) 행 제거
            rows = []
            for tl in table_lines:
                cells = [c.strip() for c in tl.strip("|").split("|")]
                if all(re.fullmatch(r":?-{2,}:?", c or "-") for c in cells):
                    continue
                rows.append(cells)
            if rows:
                ncol = max(len(r) for r in rows)
                table = doc.add_table(rows=0, cols=ncol)
                table.style = "Light Grid Accent 1"
                for ridx, cells in enumerate(rows):
                    cells = cells + [""] * (ncol - len(cells))
                    row = table.add_row().cells
                    for cidx, ctext in enumerate(cells):
                        row[cidx].text = ""
                        par = row[cidx].paragraphs[0]
                        add_runs_with_bold(par, ctext)
                        if ridx == 0:
                            for rr in par.runs:
                                rr.bold = True
            continue

        # 헤딩
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            text = strip_inline_math(m.group(2)).replace("*", "")
            if level == 1:
                h = doc.add_heading(text, level=0)
            else:
                h = doc.add_heading(text, level=min(level - 1, 4))
            i += 1
            continue

        # 인용/메모
        if stripped.startswith(">"):
            text = stripped.lstrip(">").strip()
            p = doc.add_paragraph()
            r_par = add_runs_with_bold(p, text)
            for rr in p.runs:
                rr.italic = True
                rr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            i += 1
            continue

        # 순서 목록
        m = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_runs_with_bold(p, m.group(2))
            i += 1
            continue

        # 비순서 목록
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs_with_bold(p, stripped[2:])
            i += 1
            continue

        # 일반 문단
        p = doc.add_paragraph()
        add_runs_with_bold(p, stripped)
        i += 1

    doc.save(dst)
    print(f"saved: {dst}")


if __name__ == "__main__":
    main()
